"""
InsightOS — Report Generation Engine
Generates PDF, HTML, and DOCX reports.
"""
import pandas as pd
from typing import Dict, Any
from pathlib import Path
import json
import warnings
warnings.filterwarnings("ignore")

REPORTS_DIR = Path(__file__).parent.parent / "uploads" / "reports"
REPORTS_DIR.mkdir(parents=True, exist_ok=True)


def generate_html_report(story: Dict[str, Any], health: Dict[str, Any], filename: str, dataset_id: str) -> str:
    """Generate a professional HTML report."""
    investigations_html = ""
    for inv in story.get("investigations", []):
        badge_color = "#22c55e" if inv["confidence"] >= 90 else "#f59e0b" if inv["confidence"] >= 70 else "#ef4444"
        investigations_html += f"""
        <div class="investigation-card">
          <div class="inv-header">
            <span class="inv-number">#{inv['id']}</span>
            <span class="confidence-badge" style="background:{badge_color}">{inv['confidence']}% confidence</span>
          </div>
          <h3>{inv['title']}</h3>
          <p class="explanation">{inv['explanation']}</p>
          <div class="evidence-box">
            <strong>Statistical Test:</strong> {inv['statistical_test']}<br>
            <strong>Business Impact:</strong> {inv.get('business_impact','N/A').upper()}
          </div>
        </div>"""

    key_findings_html = "".join(f"<li>{f}</li>" for f in story.get("key_findings", []))
    risks_html = "".join(f"<li class='risk'>{r}</li>" for r in story.get("risks", []))
    recs_html = "".join(f"<li class='rec'>{r}</li>" for r in story.get("recommendations", []))

    dims = health.get("dimensions", {})
    dimension_bars = ""
    for dim_name, dim_data in dims.items():
        score = dim_data.get("score", 0)
        color = "#22c55e" if score >= 80 else "#f59e0b" if score >= 60 else "#ef4444"
        dimension_bars += f"""
        <div class="dim-bar">
          <div class="dim-label">{dim_name.capitalize()}</div>
          <div class="bar-container">
            <div class="bar-fill" style="width:{score}%;background:{color}"></div>
          </div>
          <div class="dim-score">{score:.1f}%</div>
        </div>"""

    html = f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>InsightOS Report — {filename}</title>
<style>
  * {{ margin:0; padding:0; box-sizing:border-box; }}
  body {{ font-family:'Segoe UI',sans-serif; background:#0f0f1a; color:#e2e8f0; line-height:1.6; }}
  .container {{ max-width:900px; margin:0 auto; padding:40px 20px; }}
  .header {{ text-align:center; padding:40px 0; border-bottom:2px solid #6366f1; margin-bottom:40px; }}
  .header h1 {{ font-size:2.5rem; background:linear-gradient(135deg,#6366f1,#8b5cf6); -webkit-background-clip:text; -webkit-text-fill-color:transparent; }}
  .header p {{ color:#94a3b8; margin-top:8px; font-size:1.1rem; }}
  .section {{ background:#1e1e2e; border-radius:12px; padding:28px; margin-bottom:24px; border:1px solid #2d2d3d; }}
  .section h2 {{ color:#6366f1; font-size:1.4rem; margin-bottom:16px; }}
  .health-score {{ text-align:center; padding:20px; }}
  .score-circle {{ font-size:4rem; font-weight:bold; color:#6366f1; }}
  .grade {{ font-size:2rem; color:#8b5cf6; }}
  .dim-bar {{ display:flex; align-items:center; gap:12px; margin:8px 0; }}
  .dim-label {{ width:120px; font-size:0.9rem; color:#94a3b8; text-transform:capitalize; }}
  .bar-container {{ flex:1; height:8px; background:#2d2d3d; border-radius:4px; overflow:hidden; }}
  .bar-fill {{ height:100%; border-radius:4px; transition:width 0.3s; }}
  .dim-score {{ width:50px; text-align:right; font-size:0.9rem; }}
  .investigation-card {{ background:#16213e; border:1px solid #6366f1; border-radius:10px; padding:20px; margin:12px 0; }}
  .inv-header {{ display:flex; justify-content:space-between; align-items:center; margin-bottom:10px; }}
  .inv-number {{ color:#6366f1; font-weight:bold; font-size:1.1rem; }}
  .confidence-badge {{ padding:4px 12px; border-radius:20px; color:white; font-size:0.85rem; font-weight:600; }}
  .investigation-card h3 {{ color:#e2e8f0; font-size:1.1rem; margin:8px 0; }}
  .explanation {{ color:#94a3b8; margin:8px 0; }}
  .evidence-box {{ background:#0f0f1a; border-radius:6px; padding:12px; margin-top:10px; font-size:0.9rem; color:#a5b4fc; }}
  ul {{ padding-left:20px; }}
  li {{ margin:6px 0; color:#94a3b8; }}
  li.risk {{ color:#fca5a5; }}
  li.rec {{ color:#86efac; }}
  .footer {{ text-align:center; padding:20px; color:#4b5563; font-size:0.85rem; border-top:1px solid #2d2d3d; margin-top:40px; }}
  .exec-summary {{ white-space:pre-wrap; color:#94a3b8; font-size:0.95rem; }}
</style>
</head>
<body>
<div class="container">
  <div class="header">
    <h1>InsightOS</h1>
    <p>Automated Data Intelligence Report</p>
    <p style="margin-top:4px;color:#6366f1;font-weight:600;">{filename}</p>
  </div>

  <div class="section">
    <h2>📋 Executive Summary</h2>
    <div class="exec-summary">{story.get('executive_summary','').replace('## Executive Summary','').replace('**','').strip()}</div>
  </div>

  <div class="section">
    <h2>🏥 Data Health Score</h2>
    <div class="health-score">
      <div class="score-circle">{health.get('overall_score',0):.1f}</div>
      <div class="grade">Grade: {health.get('grade','?')} — {health.get('status','')}</div>
    </div>
    {dimension_bars}
    <p style="margin-top:16px;color:#94a3b8;font-style:italic;">{health.get('summary','')}</p>
  </div>

  <div class="section">
    <h2>🔍 Key Findings</h2>
    <ul>{key_findings_html}</ul>
  </div>

  <div class="section">
    <h2>🔬 Investigations</h2>
    {investigations_html if investigations_html else '<p style="color:#94a3b8">No significant patterns found.</p>'}
  </div>

  <div class="section">
    <h2>⚠️ Potential Risks</h2>
    <ul>{risks_html}</ul>
  </div>

  <div class="section">
    <h2>💡 Recommendations</h2>
    <ul>{recs_html}</ul>
  </div>

  <div class="footer">
    Generated by InsightOS — Intelligent Platform for Automated EDA and Pattern Discovery
  </div>
</div>
</body>
</html>"""

    report_path = REPORTS_DIR / f"{dataset_id}_report.html"
    with open(report_path, "w", encoding="utf-8") as f:
        f.write(html)
    return str(report_path)


def generate_pdf_report(story: Dict[str, Any], health: Dict[str, Any], filename: str, dataset_id: str) -> str:
    """Generate a PDF report using ReportLab."""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib import colors
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
        from reportlab.lib.enums import TA_CENTER, TA_LEFT

        report_path = REPORTS_DIR / f"{dataset_id}_report.pdf"
        doc = SimpleDocTemplate(str(report_path), pagesize=letter,
                                leftMargin=0.75*inch, rightMargin=0.75*inch,
                                topMargin=0.75*inch, bottomMargin=0.75*inch)
        styles = getSampleStyleSheet()
        story_elements = []

        # Title
        title_style = ParagraphStyle("title", parent=styles["Title"],
                                      textColor=colors.HexColor("#6366f1"), fontSize=24)
        story_elements.append(Paragraph("InsightOS — Intelligence Report", title_style))
        story_elements.append(Paragraph(f"Dataset: {filename}", styles["Normal"]))
        story_elements.append(Spacer(1, 0.2*inch))
        story_elements.append(HRFlowable(width="100%", color=colors.HexColor("#6366f1")))
        story_elements.append(Spacer(1, 0.2*inch))

        # Health score
        h2 = ParagraphStyle("h2", parent=styles["Heading2"], textColor=colors.HexColor("#6366f1"))
        story_elements.append(Paragraph("Data Health Score", h2))
        score = health.get("overall_score", 0)
        story_elements.append(Paragraph(
            f"Overall Score: <b>{score:.1f}/100</b> — Grade: <b>{health.get('grade','?')}</b> ({health.get('status','')})",
            styles["Normal"]
        ))
        story_elements.append(Spacer(1, 0.15*inch))

        # Key findings
        story_elements.append(Paragraph("Key Findings", h2))
        for finding in story.get("key_findings", []):
            clean = finding.replace("**", "").replace("`", "")
            story_elements.append(Paragraph(f"• {clean}", styles["Normal"]))
        story_elements.append(Spacer(1, 0.15*inch))

        # Recommendations
        story_elements.append(Paragraph("Recommendations", h2))
        for rec in story.get("recommendations", []):
            story_elements.append(Paragraph(f"• {rec}", styles["Normal"]))
        story_elements.append(Spacer(1, 0.15*inch))

        # Investigations
        story_elements.append(Paragraph("Investigations", h2))
        for inv in story.get("investigations", []):
            story_elements.append(Paragraph(
                f"<b>#{inv['id']}</b>: {inv['title']} (Confidence: {inv['confidence']}%)",
                styles["Normal"]
            ))
            story_elements.append(Paragraph(inv.get("explanation", ""), styles["Normal"]))
            story_elements.append(Spacer(1, 0.1*inch))

        doc.build(story_elements)
        return str(report_path)
    except Exception as e:
        return f"ERROR: {e}"


def generate_docx_report(story: Dict[str, Any], health: Dict[str, Any], filename: str, dataset_id: str) -> str:
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        report_path = REPORTS_DIR / f"{dataset_id}_report.docx"
        doc = Document()

        # Title
        title = doc.add_heading("InsightOS — Intelligence Report", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(f"Dataset: {filename}")
        doc.add_paragraph(f"Overall Health Score: {health.get('overall_score',0):.1f}/100 — {health.get('status','')}")
        doc.add_heading("Key Findings", level=1)
        for finding in story.get("key_findings", []):
            doc.add_paragraph(finding.replace("**","").replace("`",""), style="List Bullet")
        doc.add_heading("Investigations", level=1)
        for inv in story.get("investigations", []):
            p = doc.add_paragraph()
            p.add_run(f"#{inv['id']}: {inv['title']}").bold = True
            doc.add_paragraph(f"Confidence: {inv['confidence']}% | Impact: {inv.get('business_impact','N/A').upper()}")
            doc.add_paragraph(inv.get("explanation",""))
        doc.add_heading("Recommendations", level=1)
        for rec in story.get("recommendations", []):
            doc.add_paragraph(rec, style="List Bullet")

        doc.save(str(report_path))
        return str(report_path)
    except Exception as e:
        return f"ERROR: {e}"
